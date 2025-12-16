import os, io, tempfile, pickle
from typing import List

import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI

from langchain_core.prompts import ChatPromptTemplate
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_classic.chains import create_retrieval_chain


# -----------------------------
# Robust Text Extraction
# -----------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2, fallback to pdfminer."""
    text = ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    except Exception:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            text = pdfminer_extract(tmp_path)
        except Exception:
            text = ""
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text and table data from DOCX."""
    text_parts = []
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    doc = DocxDocument(tmp_path)

    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def extract_text_from_csv(file_bytes: bytes, max_rows: int = 50) -> str:
    """Convert CSV into summarized text form."""
    try:
        df = pd.read_csv(io.BytesIO(file_bytes))
    except UnicodeDecodeError:
        df = pd.read_csv(io.BytesIO(file_bytes), encoding="latin-1")

    sample = df.head(max_rows)
    lines = [", ".join(sample.columns)]
    for _, row in sample.iterrows():
        lines.append(", ".join(map(str, row.values)))

    return "\n".join(lines)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Read plain text with encoding fallback."""
    for enc in ("utf-8", "latin-1"):
        try:
            return file_bytes.decode(enc)
        except Exception:
            continue
    return ""


# -----------------------------
# File Loader
# -----------------------------

def load_file(file) -> List[Document]:
    """Unified file loader returning list[Document]."""
    name = file.name.lower()
    data = file.read()

    if name.endswith(".pdf"):
        text = extract_text_from_pdf(data)
    elif name.endswith(".docx"):
        text = extract_text_from_docx(data)
    elif name.endswith(".csv"):
        text = extract_text_from_csv(data)
    elif name.endswith(".txt"):
        text = extract_text_from_txt(data)
    else:
        raise ValueError(f"Unsupported file: {name}")

    return [Document(page_content=text, metadata={"source": file.name})] if text.strip() else []


# -----------------------------
# Preprocessing & Embeddings
# -----------------------------

_text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=150
)

_embeddings = None  # cached globally


def get_embeddings():
    """Lazy-load embedding model once for speed."""
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-mpnet-base-v2"
        )
    return _embeddings


def process_uploaded_files(uploaded_files):
    """Extract + chunk text from uploaded files."""
    documents = []
    for f in uploaded_files:
        documents.extend(load_file(f))
    return _text_splitter.split_documents(documents)


def create_vectorstore(docs, index_path="faiss_index"):
    """Create a FAISS vector store."""
    if not docs:
        raise ValueError("No text extracted from uploaded files.")

    embeddings = get_embeddings()
    vectorstore = FAISS.from_documents(docs, embeddings)

    os.makedirs(index_path, exist_ok=True)
    with open(f"{index_path}/index.pkl", "wb") as f:
        pickle.dump(vectorstore, f)

    return vectorstore


# -----------------------------
# RAG Chain (LangChain 1.x)
# -----------------------------

def build_rag_chain(vectorstore, k: int = 5):
    """Build retrieval-augmented generation chain (LCEL)."""

    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        temperature=0.0,
        google_api_key=os.getenv("GEMINI_API_KEY"),
    )

    prompt = ChatPromptTemplate.from_messages([
        (
            "system",
            "You are a helpful assistant. Use only the provided context to answer."
        ),
        ("human", "{input}")
    ])

    document_chain = create_stuff_documents_chain(
        llm=llm,
        prompt=prompt
    )

    retriever = vectorstore.as_retriever(search_kwargs={"k": k})

    rag_chain = create_retrieval_chain(
        retriever=retriever,
        combine_docs_chain=document_chain
    )

    return rag_chain



